# -*- coding: utf-8 -*-
"""
Created on Mon Nov 11 16:34:42 2024

@author: jveraz
"""

import os
from langchain_openai import ChatOpenAI
from langchain_experimental.graph_transformers import LLMGraphTransformer
from langchain_core.documents import Document
import networkx as nx
import json
from dotenv import load_dotenv
from PyPDF2 import PdfReader
from docx import Document as DocxDocument 
from rapidfuzz import fuzz, process

############
## OPENAI ##
############

# Load environment variables from a .env file
load_dotenv()

# Set the OpenAI API key from the .env file
api_key = os.getenv("OPENAI_API_KEY")

# model
llm = ChatOpenAI(temperature=0.2, 
                 model_name="gpt-4o")
#,
#                 model_kwargs={
#        "messages": [
#            {"role": "system", "content": """
#Eres un experto en eventos culturales andinos y patrimonio cultural inmaterial peruano. En particular,
#estás especializado en la festividad de la Mamacha Carmen en Paucartambo, Cusco, Perú. La Mamacha Carmen también es conocida como Virgen del Carmen.
#Tu tarea es analizar textos relacionados con esta festividad para construir un knowledge graph. Para esto, debes extraer tantas entidades clave como sea posible (incluyendo nombres de danzas, lugares, figuras religiosas, participantes, costumbres, objetos ceremoniales y eventos). 
#Además, identifica subentidades o descompón conceptos complejos en términos más simples, y extrae todas las relaciones relevantes (como 'en honor de', 'asociada con', 'realizada en', 'utilizada en', 'patrona de', 'parte de la tradición de', 'ubicada en', 'representada por').

#Busca todas las conexiones posibles, incluso si no son evidentes a primera vista. Sé exhaustivo al identificar entidades y relaciones, y mantén todo en español, respetando los nombres originales. Si hay términos generales (por ejemplo, 'máscaras' o 'iglesia'), inclúyelos siempre que estén conectados con la festividad.
#"""}
#        ]
#    }
#)
llm_transformer = LLMGraphTransformer(llm=llm)

##############
## ARCHIVOS ##
##############

# Función para leer todos los archivos .txt, .pdf y .docx de la carpeta "archivos"
def read_files_from_folder(folder_path):
    content = ""
    for filename in os.listdir(folder_path):
        file_path = os.path.join(folder_path, filename)
        if filename.endswith(".txt"):
            with open(file_path, "r", encoding="utf-8") as file:
                content += file.read() + "\n"
        elif filename.endswith(".pdf"):
            pdf_reader = PdfReader(file_path)
            for page in pdf_reader.pages:
                content += page.extract_text() + "\n"
        elif filename.endswith(".docx"):
            doc = DocxDocument(file_path)
            for paragraph in doc.paragraphs:
                content += paragraph.text + "\n"
    return content

# Leer todos los archivos de la carpeta "archivos"
folder_path = "archivos"
lines = read_files_from_folder(folder_path)

###################
## NORMALIZACIÓN ##
###################

# Lista de palabras o nombres a ignorar
IGNORED_WORDS = ["milagros alccacuntor farroñan", "tourism promotion plan","encuesta",
                 "variables e indicadores","tourism", "visitor", "transportation",
                 "gastronomy", "formulación de la hipótesis", "hipótesis general", "hipótesis específicas",
                 "instrumentos de recolección de datos", "recurso turístico",
                 "actividad turística", "producto turístico", "spain"]

def should_ignore(name):
    """
    Verifica si un nombre debe ser ignorado.
    Args:
        name (str): El nombre a verificar.
    Returns:
        bool: True si debe ignorarse, False en caso contrario.
    """
    name = name.lower().strip()  # Normalizar a minúsculas
    return any(ignored_word in name for ignored_word in IGNORED_WORDS)

# Diccionario de equivalencias para normalización

NAME_EQUIVALENCES = {
    "ancelmo rojas": "anselmo rojas",
    "anselmo rojas": "anselmo rojas",
    "don anselmo rojas": "anselmo rojas",
    "don alsemo rojas": "anselmo rojas",
    "alsemo rojas": "anselmo rojas",
    "anselmo": "anselmo rojas",
    "collas" : "qolla",
    "qollas" : "qolla",
    "colla" : "qolla",
    "qollas" : "qolla",
    "macana" : "inkari makana",
    "macana" : "inkari makana",
    "incari macana" : "inkari makana",
    "danza macanas" : "inkari makana",
    "danza macana" : "inkari makana",
    "inka makana" : "inkari makana",
    "danza chunchada" : "chunchada",
    "danza chunchadas" : "chunchada",
    "danza chunchada" : "chunchada",
    "chicas de la danza chunchada" : "chunchada",
    "danza amazonas" : "chunchada",
    "chunchadas" : "chunchada",
    "danza chunchada paucartambina" : "chunchada",
    "cargowasis" : "cargo wasi",
    "cargowasi" : "cargo wasi",
    "cargo wasi" : "cargo wasi",
    "mamita del rosario": "virgen del rosario",
    "virgen del rosario": "virgen del rosario",
    "virgen_de_rosario_de_paucartambo": "virgen del rosario",
    "virgen_del_rosario_de_paucartambo": "virgen del rosario",
    "virgen_de_rosario_paucartambo": "virgen del rosario",
    "fiesta_de_la_virgen_del_rosario": "virgen del rosario",
    "mamita": "virgen del rosario",
    "merienda" : "merienda paucartambina",
    "merienda paucartambina" : "merienda paucartambina",
    "carmen": "virgen del carmen"
}

def normalize_name(name):
    """
    Normaliza un nombre utilizando un diccionario de equivalencias, eliminando títulos y errores ortográficos.
    Args:
        name (str): El nombre a normalizar.
    Returns:
        str: El nombre normalizado.
    """
    # Convertir a minúsculas y quitar espacios
    name = name.lower().strip()

    # Eliminar títulos comunes
    for title in ["don", "mr.", "mrs.", "sr.", "dr."]:
        name = name.replace(title, "").strip()

    # Si está en el diccionario de equivalencias, usar el valor mapeado
    if name in NAME_EQUIVALENCES:
        return NAME_EQUIVALENCES[name]

    # Buscar coincidencias similares usando RapidFuzz
    known_names = list(NAME_EQUIVALENCES.values())
    normalized = process.extractOne(name, known_names, scorer=fuzz.ratio)
    if normalized and normalized[1] >= 80:  # Umbral de similitud
        return normalized[0]

    # Si no hay coincidencias, devolver el nombre procesado
    return name

########################
## DIVISIÓN EN CHUNKS ##
########################

# Dividir el texto en fragmentos de 1000 caracteres (ajusta este tamaño si es necesario)
chunk_size = 10000
text_chunks = [lines[i:i + chunk_size] for i in range(0, len(lines), chunk_size)]

# Crear grafo de NetworkX
nx_graph = nx.Graph()

# Procesar cada fragmento individualmente
for i, chunk in enumerate(text_chunks):
    print(f"Procesando fragmento {i + 1} de {len(text_chunks)}...")

    # Crear un documento con el fragmento de texto
    document = Document(page_content=chunk)
  
    graph_documents = llm_transformer.convert_to_graph_documents([document])
    
    # Agregar nodos y aristas del fragmento al grafo
    for doc in graph_documents:
        for node in doc.nodes:
            node_id = normalize_name(dict(node)["id"])
            content =normalize_name( dict(node)["type"])
            
            # Ignorar nodos que coincidan con palabras específicas
            if should_ignore(node_id) or should_ignore(content):
                continue
            
            nx_graph.add_node(node_id, content=content)
        
        for edge in doc.relationships:
            edge_dict = vars(edge) if not isinstance(edge, dict) else edge
            source = normalize_name(str(dict(edge_dict["source"])["id"]))
            target = normalize_name(str(dict(edge_dict["target"])["id"]))
            tipo = edge_dict["type"]
            
            # Ignorar relaciones que incluyan palabras específicas
            if should_ignore(source) or should_ignore(target):
                continue
            
            nx_graph.add_edge(source, target, content=tipo)
            
# Guardar el grafo en JSON
data = nx.node_link_data(nx_graph)
with open("grafo/graph.json", "w") as file:
    json.dump(data, file)


